import os
import docx
from langchain import PromptTemplate, LLMChain
from langchain.chains.question_answering import load_qa_chain

from settings import openai_api_key
from langchain.llms import OpenAI
import streamlit as st
from langchain.document_loaders import Docx2txtLoader
from docx import Document
from langchain.output_parsers import RegexParser

os.environ['OPENAI_API_KEY'] = openai_api_key

llm = OpenAI(temperature=1, verbose=True)

loader = Docx2txtLoader('MSAReviewPlaybook.docx')
pages = loader.load_and_split()
playbook_docs = loader.load_and_split()

template = """{context}
List any problems with this MSA section:
{msa_section}
"""

# r"Issues: (.*?)\nReason: (.*?)\nScore: (.*)"

output_parser = RegexParser(
    regex=r"(.*?)\nReason: (.*?)\nScore: (.*)",
    output_keys=["answer", "reason", "score"],
)

prompt_template = """Use the following pieces of context to identify unacceptable language with a section of an MSA, and why it is unacceptable.
In addition to listing the issue and the reason, also return a score of you confident you are that issues are present.
If no issue found, answer "None" for issues and reason, and 0 for score.
Use the exact format below to answer:

Issues: [answer here, "None" if none]
Reason: [reason for answer, "None" if none]
Score: [score between 0 and 100]

Begin!

Context:
---------
{context}
---------
MSA Section: {msa_section}
Any unacceptable language found, and why it is unacceptable:"""
PROMPT = PromptTemplate(
    template=prompt_template,
    input_variables=["context", "msa_section"],
    output_parser=output_parser,
)

chain = load_qa_chain(llm, chain_type="map_rerank", return_intermediate_steps=True, prompt=PROMPT)

# prompt = PromptTemplate(
#     input_variables=["context", "msa_section"], template=template
# )

# memory = ConversationBufferMemory(memory_key="chat_history", input_key="msa_draft")
# chain = load_qa_chain(
#     llm, chain_type="map_rerank", prompt=prompt
# )

st.title('🦜🔗 MSA Review Playbook')
# Create a text input box for the user
# uploaded_file = st.file_uploader('Submit MSA Draft (docx)', type=['docx'])
section = st.text_input("MSA Section Text")

if section:
    try:
        # doc: Document = Document(uploaded_file)
        #
        # content = '\n'.join([p.text for p in docx.Document(uploaded_file).paragraphs])
        # chunk_size = 100
        # prompt_chunks = [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]
        #
        # for chunk in prompt_chunks:
        #     chain({"input_documents": playbook_doc, "msa_draft": chunk}, return_only_outputs=True)

        answer = chain({"input_documents": playbook_docs, "msa_section": section}, return_only_outputs=True)

        # for step in answer["intermediate_steps"]:
        #     if step["score"] != "0":
        #         st.write(f'Issue: {step["answer"]}\nReason: {step["reason"]}')

        st.write(answer)

    except Exception as e:
        st.write(e)

